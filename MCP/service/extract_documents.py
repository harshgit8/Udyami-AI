"""
Standalone script to extract text from documents
"""
import json
import logging
from pathlib import Path
from datetime import datetime
from typing import Optional

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

UPLOADS_BASE_DIR = Path("./uploads")
LOCAL_USER = "local_user"


def extract_pdf(file_path: Path) -> Optional[str]:
    """Extract text from PDF file"""
    try:
        from pypdf import PdfReader
        
        reader = PdfReader(str(file_path))
        text_parts = []
        
        for page_num, page in enumerate(reader.pages, 1):
            try:
                text = page.extract_text()
                if text.strip():
                    text_parts.append(f"--- Page {page_num} ---\n{text}")
            except Exception as e:
                logger.warning(f"Error extracting page {page_num}: {e}")
                continue
        
        full_text = "\n\n".join(text_parts)
        logger.info(f"  Extracted {len(full_text)} characters from {len(reader.pages)} pages")
        return full_text
    except Exception as e:
        logger.error(f"  PDF extraction error: {e}")
        return None


def extract_docx(file_path: Path) -> Optional[str]:
    """Extract text from DOCX file"""
    try:
        from docx import Document
        
        doc = Document(str(file_path))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        full_text = "\n\n".join(text_parts)
        logger.info(f"  Extracted {len(full_text)} characters")
        return full_text
    except Exception as e:
        logger.error(f"  DOCX extraction error: {e}")
        return None


def extract_text_file(file_path: Path) -> Optional[str]:
    """Extract text from plain text files"""
    try:
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                logger.info(f"  Extracted {len(text)} characters using {encoding}")
                return text
            except UnicodeDecodeError:
                continue
        logger.error(f"  Could not decode text file")
        return None
    except Exception as e:
        logger.error(f"  Text file extraction error: {e}")
        return None


def extract_text(file_path: Path) -> Optional[str]:
    """Extract text from a document file"""
    ext = file_path.suffix.lower()
    
    if ext == '.pdf':
        return extract_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return extract_docx(file_path)
    elif ext in ['.txt', '.md', '.csv', '.json', '.xml', '.html']:
        return extract_text_file(file_path)
    else:
        logger.warning(f"  Unsupported file type: {ext}")
        return None


def create_summary(text: str, max_length: int = 300) -> str:
    """Create a simple summary from extracted text"""
    if not text:
        return "No content extracted"
    
    text = " ".join(text.split())
    
    if len(text) <= max_length:
        return text
    
    return text[:max_length].rsplit(' ', 1)[0] + "..."


def reprocess_documents():
    """Reprocess all existing documents to extract text content"""
    
    base_dir = UPLOADS_BASE_DIR / LOCAL_USER
    original_dir = base_dir / "original"
    processed_dir = base_dir / "processed"
    
    if not original_dir.exists():
        logger.error(f"Original directory not found: {original_dir}")
        return
    
    processed_count = 0
    error_count = 0
    
    for category in ["docs", "media", "links"]:
        category_original = original_dir / category
        category_processed = processed_dir / category
        
        if not category_original.exists():
            logger.info(f"Skipping category {category} - directory not found")
            continue
        
        category_processed.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"\nProcessing category: {category}")
        logger.info("=" * 60)
        
        for file_path in category_original.iterdir():
            if file_path.is_file():
                try:
                    logger.info(f"\nProcessing: {file_path.name}")
                    
                    # Extract text content
                    extracted_text = None
                    if category == "docs":
                        extracted_text = extract_text(file_path)
                        if extracted_text:
                            logger.info(f"  ✓ Extracted {len(extracted_text)} characters")
                        else:
                            logger.warning(f"  ⚠ No text extracted")
                    
                    file_stem = file_path.stem
                    file_ext = file_path.suffix
                    
                    # Create summary
                    if extracted_text:
                        summary = create_summary(extracted_text, max_length=300)
                    else:
                        summary = f"Local file: {file_path.name}"
                    
                    # Update metadata
                    metadata = {
                        "old_name": file_path.name,
                        "name": file_stem,
                        "summary": summary,
                        "tags": ["local", category, "reprocessed"],
                        "created_at": datetime.now().isoformat(),
                        "has_content": extracted_text is not None,
                        "content_length": len(extracted_text) if extracted_text else 0
                    }
                    
                    meta_path = category_processed / f"{file_stem}.meta"
                    with open(meta_path, "w", encoding="utf-8") as f:
                        json.dump(metadata, f, indent=2, ensure_ascii=False)
                    logger.info(f"  ✓ Updated metadata")
                    
                    # Create/update markdown with full content
                    if extracted_text:
                        md_content = f"""# {file_path.name}

**File:** {file_stem}{file_ext}
**Category:** {category}
**Reprocessed:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
**Content Length:** {len(extracted_text)} characters

---

## Extracted Content

{extracted_text}
"""
                    else:
                        md_content = f"""# {file_path.name}

**File:** {file_stem}{file_ext}
**Category:** {category}
**Reprocessed:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

File uploaded locally. Text extraction not available for this file type.
"""
                    
                    md_path = category_processed / f"{file_stem}.md"
                    with open(md_path, "w", encoding="utf-8") as f:
                        f.write(md_content)
                    logger.info(f"  ✓ Created markdown")
                    
                    processed_count += 1
                    
                except Exception as e:
                    logger.error(f"  ✗ Error: {e}")
                    error_count += 1
                    continue
    
    logger.info("\n" + "=" * 60)
    logger.info(f"Reprocessing complete!")
    logger.info(f"  Successfully processed: {processed_count} files")
    logger.info(f"  Errors: {error_count} files")
    logger.info("=" * 60)


if __name__ == "__main__":
    logger.info("Starting document extraction...")
    logger.info(f"Base directory: {UPLOADS_BASE_DIR.absolute()}\n")
    reprocess_documents()
